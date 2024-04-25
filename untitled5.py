import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
import pandas as pd
import openpyxl
from openpyxl.styles import Font
from openpyxl.styles import Alignment
from datetime import datetime
from docx import Document
import os
import re
from docx.shared import Pt
import locale


try:
    locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')
except locale.Error:
    print("Локаль не поддерживается, используется локаль по умолчанию.")

desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
main_folder_path = os.path.join(desktop_path, 'Document_Sets')
if not os.path.exists(main_folder_path):
    os.makedirs(main_folder_path)

db_path = "/Users/maksimpekerman/Desktop/Full_School_Table.xlsx"

excel_invoice_path_pekerman = "/Users/maksimpekerman/Desktop/Счет №.xlsx"
excel_act_path_pekerman = "/Users/maksimpekerman/Desktop/АКТ №.xlsx"
word_contract_path_pekerman = "/Users/maksimpekerman/Desktop/контракт №.docx"

excel_invoice_path_gredushko = "/Users/maksimpekerman/Desktop/Счет № Гредюшко.xlsx"
excel_act_path_gredushko = "/Users/maksimpekerman/Desktop/АКТ № Гредюшко.xlsx"
word_contract_path_gredushko = "/Users/maksimpekerman/Desktop/Контракт № Гредюшко.docx"

service_dict = {
        'Неисправность сетей ГВС': 'Ремонт системы ГВС',
        'Неисправность сетей канализации': 'Прочистка канализации',
        'Неисправность сетей ХВС': 'Ремонт системы ХВС',
        'Подвальное помещение и техническое подполье':'Прочистка канализации',
        'Горячее водоснабжение': 'Ремонт системы ГВС',
        'Крыша и водосточная система': 'Прочистка канализации'
    }

TNR_12 = Font(name='Times New Roman', size=12, bold=False)
TNR_12_bold = Font(name='Times New Roman', size=12, bold=True)
    
arial_12 = Font(name='Arial', size = 12, bold=False)
arial_12_bold = Font(name='Arial', size = 12, bold=True)
    
alignment_center = Alignment(horizontal='center', vertical='center')
    
def main_app_window():
    root = tk.Tk()
    root.title("Обработка и обновление документов")
    tk.Button(root, text="Загрузить файл пользователя", command=lambda: update_files(load_user_file(), root)).pack(pady=20)
    root.mainloop()

def load_user_file():
    load_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if load_file_path:
        print(f"Загружен файл: {load_file_path}")
        return pd.read_excel(load_file_path)
    else:
        messagebox.showwarning("Предупреждение", "Файл не выбран.")
        return None

def load_database():
    try:
        db = pd.read_excel(db_path)
        print("База данных успешно загружена.")
        return db
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось загрузить базу данных: {str(e)}")
        return None

def request_initial_number(root):
    initial_number = simpledialog.askinteger("Input", "Введите начальный номер для документов:", parent=root)
    if initial_number is not None:
        print(f"Выбран начальный номер: {initial_number}")
    else:
        messagebox.showinfo("Информация", "Начальный номер не введен.")
    return initial_number

def formated_principal_name(full_name):
    parts = full_name.split()
    if len(parts) > 1:  # Проверяем, есть ли в имени более одного компонента
        # Возвращаем имя в формате 'Фамилия И.О.'
        return f"{parts[0]} {'.'.join([p[0] for p in parts[1:]])}."
    else:
        # Если имя состоит только из одной части, возвращаем как есть
        return full_name
    
def declension_surname(surname, gender):
    
    if gender == "M":  # Для мужского пола
        if surname.endswith("ий") or surname.endswith("ый"):
            return surname[:-2] + "ого"
        elif surname.endswith("ых"):
            return surname
        else:
            return surname + "a"
        
    elif gender == "F":  # Для женского пола
        if surname.endswith("ая"):
            return surname[:-2] + "ой"
        elif surname.endswith("а"):
            return surname[:-1] + "ой"
    return surname

def decline_action_word(gender):
    if gender == "M":
        return "действующего"
    elif gender == "F":
        return "действующей"
    return "действующе"

def update_invoice(service_dict,worksheet, row, db_row, initial_number, service_type):
    try:    
        if "Пекерман" in str(row[16]):
            
            worksheet['D12'].font = TNR_12_bold
            
            worksheet['B12'].value = f"Счет на оплату № {initial_number} от"
            worksheet['B12'].font = TNR_12_bold
            
            worksheet['D14'].value = db_row['Full_School_Name']
            worksheet['D14'].font = TNR_12
            
            worksheet['B15'].value = f"Основание: Контракт № {initial_number} от"
            worksheet['B15'].font = TNR_12
            
            worksheet['D15'].font = TNR_12
    
            worksheet['C18'].value = f"За выполненные работы по аварийному ремонту в здании\n{db_row['Short_School_Name']}"
            worksheet['C18'].font = TNR_12
            worksheet['C18'].alignment = alignment_center
    
            if service_type in service_dict:
                worksheet['C19'].value = service_dict[service_type]
                worksheet['C19'].font = TNR_12
                worksheet['C19'].alignment = alignment_center
                
        if "Гредюшко" in str(row[16]):
            
            worksheet['B12'].value = f"Счет на оплату № {initial_number} от"
            worksheet['B12'].font = arial_12_bold
            
            worksheet['D12'].font = arial_12_bold
            
            worksheet['C14'].value = db_row['Full_School_Name']
            worksheet['C14'].font = arial_12
            
            worksheet['C17'].value = f"За выполненные аварийные работы  по контракту № {initial_number}"
            worksheet['C17'].font = arial_12
            worksheet['C17'].alignment = alignment_center
            
            if service_type in service_dict:
                worksheet['B18'].value = service_dict[service_type]
                worksheet['B18'].font = arial_12
                worksheet['B18'].alignment = alignment_center
    except Exception as e:
        messagebox.showerror("Ошибка счет", f"Не удалось обновить счет файл: {str(e)}")

def update_date_in_second_paragraph(doc):
    if len(doc.paragraphs) > 1:
        paragraph = doc.paragraphs[1]  # Предполагаем, что дата находится во втором параграфе
        old_date_pattern = r'\d{1,2} \w+ \d{4} г\.'  # Регулярное выражение для поиска даты
        
        new_date = datetime.now().strftime('%d %B %Y г.')  # Новая дата в нужном формате на русском
        
        # Изменяем текст в каждом run параграфа
        for run in paragraph.runs:
            if re.search(old_date_pattern, run.text):
                run.text = re.sub(old_date_pattern, new_date, run.text)
                # Устанавливаем шрифт и размер шрифта
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = False
            
def update_third_paragraph(doc, db_row):
    if len(doc.paragraphs) >= 3:  # Убедимся, что есть хотя бы три параграфа
        # Получаем полный текст третьего параграфа
        paragraph = doc.paragraphs[2]
        
        # Форматируем имя директора и слово "действующий/ая" в зависимости от пола
        principal_name = formated_principal_name(db_row["Principal_Name"])
        surname_declined = declension_surname(principal_name.split()[0], db_row["Sex"])
        full_principal_name = surname_declined + ' ' + ' '.join(principal_name.split()[1:])
        action_word = decline_action_word(db_row["Sex"])
        
        if re.search(r"Детский", db_row['Full_School_Name'], re.IGNORECASE):
            Principal_title = "заведующей"
        else:
            Principal_title = "директора"
        # Формируем новый текст параграфа
        new_text = f"{db_row['Full_School_Name']}, именуемое в дальнейшем «Заказчик» в лице {Principal_title} {full_principal_name}, {action_word} на основании Устава с одной стороны, и Индивидуальный Предприниматель Пекерман Леонид Ильич, именуемый в дальнейшем «Подрядчик», в соответствии с п.4 ч.1 ст.93 ФЗ от 05.04.2013 № 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, услуг для обеспечения государственных и муниципальных нужд», заключили настоящий Контракт о нижеследующем:"
        
        # Обновляем текст в параграфе
        paragraph.text = new_text                
        
        paragraph.clear()
        
        # Создаём новый run с изменённым текстом
        run = paragraph.add_run(new_text)
        # Настраиваем шрифт и размер шрифта
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = False
        
def update_forth_paragraph(doc):
    if len(doc.paragraphs) >= 7:
        paragraph = doc.paragraphs[6]
        current_date = datetime.now().strftime('%d.%m.%Y')
        new_text = f'начало –{current_date} года;   окончание –{current_date} года.'
        
        paragraph.text = new_text                
        
        paragraph.clear()
        
        # Создаём новый run с изменённым текстом
        run = paragraph.add_run(new_text)
        # Настраиваем шрифт и размер шрифта
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = False
        
def update_word_table(doc, db_row):
    table = doc.tables[0]
    cell = table.cell(1, 0)
    
    # Удаляем текст из всех существующих runs в первом абзаце ячейки
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""

    # Если в ячейке нет абзацев, добавляем новый
    if len(cell.paragraphs) == 0:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]

    # Форматирование и добавление текста
    district_modified = db_row['District'][:-2] + "ого района г.Казани"
    address_with_district = f"{db_row['Short_School_Name']} {district_modified}"
    run = p.add_run(f"{address_with_district}\n")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    text_elements = [
        (f"Адрес {db_row['Address']}\n", 'Times New Roman', 11),
        (f"ИНН {db_row['INN']}\n", 'Times New Roman', 11),
        (f"КПП {db_row['KPP']}\n", 'Times New Roman', 11),
        (f"р\\с 03234643927010001100 Отделение – НБ Республика Татарстан Банка России/УФК по Республике Татарстан г. Казань \n", 'Times New Roman', 11),
        (f"БИК 019205400\n", 'Times New Roman', 11),
        (f"Кор. сч. 40102810445370000079", 'Times New Roman', 11)
    ]
    for text, font_name, font_size in text_elements:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)

def insert_name_in_document(doc, full_name):
    if len(doc.paragraphs) >= 50:
        paragraph = doc.paragraphs[49]  # Индексация начинается с 0, поэтому 50-й параграф будет под индексом 49
        if len(paragraph.runs) >= 2:
            formatted_name = formated_principal_name(full_name)
            run_to_edit = paragraph.runs[0]
            run_to_edit.text = run_to_edit.text.replace('/ /', f'/{formatted_name}/')
        else:
            print("Недостаточно Run объектов в параграфе для вставки имени.")
    else:
        print("Документ не содержит достаточно параграфов для вставки имени.")

        
        
def update_word_documents(doc, row, db_row, initial_number):
    try:
        if "Пекерман" in str(row[16]):
            for paragraph in doc.paragraphs:
                if "КОНТРАКТ №" in paragraph.text:
                    for run in paragraph.runs:  # Итерация по всем runs внутри параграфа
                        if "КОНТРАКТ №" in run.text:
                            # Заменяем текст в run
                            run.text = run.text.replace("КОНТРАКТ №", f"КОНТРАКТ № {initial_number}")
                            # Установка шрифта и размера
                            run.font.name = 'Times New Roman'  # Установка шрифта
                            run.font.size = Pt(12)  # Установка размера шрифта в 14 пунктов
                            run.bold = True  # Делаем текст жирным
                    break
        elif "Гредюшко" in str(row[16]):
            for paragraph in doc.paragraphs:
                if "№" in paragraph.text[0]:
                    for run in paragraph.runs:
                        if "№" in run.text:
                            run.text = run.text.replace("№", f"№ {initial_number}")
                            # Установка шрифта и размера
                            run.font.name = 'Times New Roman'  # Установка шрифта
                            run.font.size = Pt(11)  # Установка размера шрифта в 14 пунктов
                            run.bold = True  # Делаем текст жирным
                    break
        
        update_date_in_second_paragraph(doc)
        
        update_third_paragraph(doc, db_row) 
        
        update_forth_paragraph(doc)
        
        update_word_table(doc, db_row)
        
        insert_name_in_document(doc, db_row["Principal_Name"])
    except Exception as e:
        messagebox.showerror("Ошибка word", f"Не удалось обновить word файлы: {str(e)}")

def update_act(service_dict,worksheet, db_row, row, initial_number, service_type):
    try:
        worksheet['A5'].value = f"Заказчик:  {db_row['Short_School_Name']}"
        worksheet['A13'].value = initial_number
        formatted_name = formated_principal_name(db_row['Principal_Name'])
        if "Пекерман" in str(row[16]):
                
            worksheet['I8'].value = initial_number
           
            worksheet['E16'].value = initial_number
    
            worksheet['C22'].value = row.iloc[13]
        
            worksheet['B29'].value = f"Директор {formatted_name}"
            
            if service_type in service_dict:
                worksheet['C24'].value = service_dict[service_type]
                worksheet['C24'].font = TNR_12
                worksheet['C24'].alignment = alignment_center
        
            worksheet['C24'].font = Font(name='Times New Roman', size=10, bold=True)
            worksheet['C24'].alignment = Alignment(horizontal='left')
        
        elif "Гредюшко" in str(row[16]):
        
            worksheet['J8'].value = initial_number
            
            worksheet['C23'].value = row.iloc[13]
            
            if service_type in service_dict:
                worksheet['C26'].value = service_dict[service_type]
                worksheet['C26'].font = arial_12
                worksheet['C26'].alignment = alignment_center
                
            worksheet['B29'].value = f"Директор {formatted_name}"
    except Exception as e:
        messagebox.showerror("Ошибка акт", f"Не удалось обновить акт файл: {str(e)}")
def update_files(data_df, root): 
    if data_df is None:
        return
    db_df = load_database()
    if db_df is None:
        return
    initial_number = request_initial_number(root)
    if initial_number is None:
        messagebox.showinfo("Информация", "Начальный номер не введен.")
        return

    try:   
        for index, row in data_df.iterrows():
            try:
                service_type = row[14]  # Assuming row[14] contains the service type
                if service_type is None or service_type not in service_dict:
                    print(f"Invalid or missing service type for row {index}.")
                continue
                if row[11] == 5088 and 'Итого' not in str(row[1]):
                    try:
                        if pd.notna(row[2]):
                            pattern = re.compile(f"({re.escape(str(row[1]))}.*\\b{re.escape(str(row[2]))}\\b)|" f"(\\b{re.escape(str(row[2]))}\\b.*{re.escape(str(row[1]))})", re.IGNORECASE)
                        else:
                            pattern = re.compile(re.escape(str(row[1])), re.IGNORECASE)
                    except Exception as e:
                        print(f"Ошибка при создании регулярного выражения: {e}")
                        continue
    
                    for db_index, db_row in db_df.iterrows():
                        try:
                            search_target = str(db_row["Short_School_Name"])
                            if "СОШ" in search_target or "ООШ" in search_target:
                                search_target = str(db_row["Full_School_Name"])
                        except Exception as e:
                            print(f"Ошибка при формировании строки поиска: {e}")
                            continue
    
                        try:
                            if pattern.search(search_target):
                                set_folder_path = os.path.join(main_folder_path, f"Set_{initial_number}")
                                os.makedirs(set_folder_path, exist_ok=True)
                        except Exception as e:
                            print(f"Ошибка при создании директории или проверке соответствия паттерну: {e}")
                            continue
    
                        try:
                            if "Пекерман" in str(row[16]):
                                workbook_invoice = openpyxl.load_workbook(excel_invoice_path_pekerman)
                                worksheet_invoice = workbook_invoice.active
                                workbook_act = openpyxl.load_workbook(excel_act_path_pekerman)
                                worksheet_act = workbook_act.active
                                doc_contract = Document(word_contract_path_pekerman)
                                
                                update_word_documents(doc_contract, row, db_row, initial_number)
                                update_invoice(service_dict, worksheet_invoice, row, db_row, initial_number, service_type)
                                update_act(service_dict, worksheet_act, row, db_row, initial_number, service_type)
                
                                workbook_invoice.save(os.path.join(set_folder_path, f"Счет № {initial_number}.xlsx"))
                                workbook_act.save(os.path.join(set_folder_path, f"АКТ № {initial_number}.xlsx"))
                                doc_contract.save(os.path.join(set_folder_path, f"Контракт № {initial_number}.docx"))
                                initial_number += 1
                            elif "Гредюшко" in str(row[16]):
                                workbook_invoice = openpyxl.load_workbook(excel_invoice_path_gredushko)
                                worksheet_invoice = workbook_invoice.active
                                workbook_act = openpyxl.load_workbook(excel_act_path_gredushko)
                                worksheet_act = workbook_act.active
                                doc_contract = Document(word_contract_path_gredushko)
                                
                                update_word_documents(doc_contract, row, db_row, initial_number)
                                update_invoice(service_dict, worksheet_invoice, row, db_row, initial_number, service_type)
                                update_act(service_dict, worksheet_act, row, db_row, initial_number, service_type)
                
                                workbook_invoice.save(os.path.join(set_folder_path, f"Счет № {initial_number}.xlsx"))
                                workbook_act.save(os.path.join(set_folder_path, f"АКТ № {initial_number}.xlsx"))
                                doc_contract.save(os.path.join(set_folder_path, f"Контракт № {initial_number}.docx"))
                                initial_number += 1
                        except Exception as e:
                            print(f"Ошибка при обработке и сохранении документов: {e}")
                        break  # Этот break должен быть внутри блока if pattern.search(search_target)
                else:
                    print(f"Строка {index} пропущена по условиям фильтрации.")
            except Exception as e:
                print(f"Ошибка при обработке строки с индексом {index}: {e}")
    except Exception as e:
        print(f"Основная ошибка во внешнем цикле: {e}")


if __name__ == "__main__":
    main_app_window()
